from fastapi import FastAPI, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, EmailStr
from typing import List, Optional
import os
import base64
from datetime import datetime, timedelta, date, time
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import tempfile
from dotenv import load_dotenv
from jose import JWTError, jwt
from passlib.context import CryptContext
from sqlalchemy import create_engine, Column, Integer, String, Date, Time, Boolean, ForeignKey
from sqlalchemy.orm import sessionmaker, declarative_base, relationship, Session
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm

load_dotenv()

app = FastAPI()

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =====================================
# Auth & Database Setup
# =====================================

DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./logbook.db")
SECRET_KEY = os.getenv("JWT_SECRET", "change-this-secret")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("JWT_EXPIRE_MINUTES", "60"))

engine = create_engine(
    DATABASE_URL, connect_args={"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {}
)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/auth/login")


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    email = Column(String, unique=True, index=True, nullable=False)
    password_hash = Column(String, nullable=False)
    is_admin = Column(Boolean, default=False, nullable=False)
    created_at = Column(String, default=lambda: datetime.utcnow().isoformat(), nullable=False)
    entries = relationship("LogbookEntryORM", back_populates="user", cascade="all, delete-orphan")


class LogbookEntryORM(Base):
    __tablename__ = "logbook_entries"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False, index=True)
    tanggal = Column(Date, nullable=False)
    jam_mulai = Column(Time, nullable=False)
    jam_selesai = Column(Time, nullable=True)
    judul_kegiatan = Column(String, nullable=False)
    rincian_kegiatan = Column(String, nullable=False)
    dokumen_pendukung = Column(String, nullable=True)  # base64 image string or URL
    user = relationship("User", back_populates="entries")


Base.metadata.create_all(bind=engine)


# =====================================
# Schemas
# =====================================

class Token(BaseModel):
    access_token: str
    token_type: str


class TokenData(BaseModel):
    user_id: Optional[int] = None


class UserCreate(BaseModel):
    email: EmailStr
    password: str
    is_admin: Optional[bool] = False


class UserOut(BaseModel):
    id: int
    email: EmailStr
    is_admin: bool

    class Config:
        from_attributes = True


class LogbookEntryIn(BaseModel):
    tanggal: date
    jam_mulai: str
    jam_selesai: Optional[str] = None
    judul_kegiatan: str
    rincian_kegiatan: str
    dokumen_pendukung: Optional[str] = None


class LogbookEntryOut(BaseModel):
    id: int
    tanggal: date
    jam_mulai: str
    jam_selesai: Optional[str]
    judul_kegiatan: str
    rincian_kegiatan: str
    dokumen_pendukung: Optional[str]

    class Config:
        from_attributes = True


# For backward compatibility with generate endpoint body
class GenerateLogbookEntry(BaseModel):
    id: str
    tanggal: str
    jam: str
    judul_kegiatan: str
    rincian_kegiatan: str
    dokumen_pendukung: Optional[str] = None


class GenerateLogbookData(BaseModel):
    entries: List[GenerateLogbookEntry]


# =====================================
# Auth Utilities
# =====================================

def verify_password(plain_password: str, password_hash: str) -> bool:
    return pwd_context.verify(plain_password, password_hash)


def get_password_hash(password: str) -> str:
    return pwd_context.hash(password)


def create_access_token(data: dict, expires_delta: Optional[timedelta] = None) -> str:
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt


def get_user_by_id(db: Session, user_id: int) -> Optional[User]:
    return db.query(User).filter(User.id == user_id).first()


def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)) -> User:
    credentials_exception = HTTPException(status_code=401, detail="Could not validate credentials")
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        uid: int = int(payload.get("sub"))
    except Exception:
        raise credentials_exception
    user = get_user_by_id(db, uid)
    if user is None:
        raise credentials_exception
    return user


def get_current_user_optional(token: Optional[str] = Depends(oauth2_scheme), db: Session = Depends(get_db)) -> Optional[User]:
    try:
        if not token:
            return None
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        uid: int = int(payload.get("sub"))
        return get_user_by_id(db, uid)
    except Exception:
        return None


# =====================================
# Auth Routes
# =====================================

@app.post("/auth/register", response_model=UserOut)
def register_user(user_in: UserCreate, db: Session = Depends(get_db)):
    existing = db.query(User).filter(User.email == user_in.email).first()
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    user = User(
        email=user_in.email,
        password_hash=get_password_hash(user_in.password),
        is_admin=bool(user_in.is_admin or False),
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return user


@app.post("/auth/login", response_model=Token)
def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == form_data.username).first()
    if not user or not verify_password(form_data.password, user.password_hash):
        raise HTTPException(status_code=400, detail="Incorrect email or password")
    access_token = create_access_token({"sub": str(user.id)})
    return {"access_token": access_token, "token_type": "bearer"}


# =====================================
# Logbook CRUD Routes (per-user, admin can see all)
# =====================================

@app.get("/logbook", response_model=List[LogbookEntryOut])
def list_logbook_entries(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    q = db.query(LogbookEntryORM)
    if not current_user.is_admin:
        q = q.filter(LogbookEntryORM.user_id == current_user.id)
    entries = q.order_by(LogbookEntryORM.tanggal.asc(), LogbookEntryORM.jam_mulai.asc()).all()
    return [
        LogbookEntryOut(
            id=e.id,
            tanggal=e.tanggal,
            jam_mulai=e.jam_mulai.strftime("%H:%M"),
            jam_selesai=e.jam_selesai.strftime("%H:%M") if e.jam_selesai else None,
            judul_kegiatan=e.judul_kegiatan,
            rincian_kegiatan=e.rincian_kegiatan,
            dokumen_pendukung=e.dokumen_pendukung,
        )
        for e in entries
    ]


@app.post("/logbook", response_model=LogbookEntryOut)
def create_logbook_entry(entry: LogbookEntryIn, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    def parse_time(value: Optional[str]) -> Optional[time]:
        if not value:
            return None
        try:
            hh, mm = value.split(":")
            return time(hour=int(hh), minute=int(mm))
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid time format, expected HH:MM")

    orm = LogbookEntryORM(
        user_id=current_user.id,
        tanggal=entry.tanggal,
        jam_mulai=parse_time(entry.jam_mulai),
        jam_selesai=parse_time(entry.jam_selesai),
        judul_kegiatan=entry.judul_kegiatan,
        rincian_kegiatan=entry.rincian_kegiatan,
        dokumen_pendukung=entry.dokumen_pendukung,
    )
    db.add(orm)
    db.commit()
    db.refresh(orm)
    return LogbookEntryOut(
        id=orm.id,
        tanggal=orm.tanggal,
        jam_mulai=orm.jam_mulai.strftime("%H:%M"),
        jam_selesai=orm.jam_selesai.strftime("%H:%M") if orm.jam_selesai else None,
        judul_kegiatan=orm.judul_kegiatan,
        rincian_kegiatan=orm.rincian_kegiatan,
        dokumen_pendukung=orm.dokumen_pendukung,
    )


@app.put("/logbook/{entry_id}", response_model=LogbookEntryOut)
def update_logbook_entry(entry_id: int, entry: LogbookEntryIn, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    orm = db.query(LogbookEntryORM).filter(LogbookEntryORM.id == entry_id).first()
    if not orm or (not current_user.is_admin and orm.user_id != current_user.id):
        raise HTTPException(status_code=404, detail="Entry not found")

    def parse_time(value: Optional[str]) -> Optional[time]:
        if not value:
            return None
        try:
            hh, mm = value.split(":")
            return time(hour=int(hh), minute=int(mm))
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid time format, expected HH:MM")

    orm.tanggal = entry.tanggal
    orm.jam_mulai = parse_time(entry.jam_mulai)
    orm.jam_selesai = parse_time(entry.jam_selesai)
    orm.judul_kegiatan = entry.judul_kegiatan
    orm.rincian_kegiatan = entry.rincian_kegiatan
    orm.dokumen_pendukung = entry.dokumen_pendukung
    db.commit()
    db.refresh(orm)
    return LogbookEntryOut(
        id=orm.id,
        tanggal=orm.tanggal,
        jam_mulai=orm.jam_mulai.strftime("%H:%M"),
        jam_selesai=orm.jam_selesai.strftime("%H:%M") if orm.jam_selesai else None,
        judul_kegiatan=orm.judul_kegiatan,
        rincian_kegiatan=orm.rincian_kegiatan,
        dokumen_pendukung=orm.dokumen_pendukung,
    )


@app.delete("/logbook/{entry_id}")
def delete_logbook_entry(entry_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    orm = db.query(LogbookEntryORM).filter(LogbookEntryORM.id == entry_id).first()
    if not orm or (not current_user.is_admin and orm.user_id != current_user.id):
        raise HTTPException(status_code=404, detail="Entry not found")
    db.delete(orm)
    db.commit()
    return {"status": "deleted"}


@app.get("/admin/logbook", response_model=List[LogbookEntryOut])
def admin_list_all_entries(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if not current_user.is_admin:
        raise HTTPException(status_code=403, detail="Admin only")
    entries = db.query(LogbookEntryORM).order_by(LogbookEntryORM.tanggal.asc(), LogbookEntryORM.jam_mulai.asc()).all()
    return [
        LogbookEntryOut(
            id=e.id,
            tanggal=e.tanggal,
            jam_mulai=e.jam_mulai.strftime("%H:%M"),
            jam_selesai=e.jam_selesai.strftime("%H:%M") if e.jam_selesai else None,
            judul_kegiatan=e.judul_kegiatan,
            rincian_kegiatan=e.rincian_kegiatan,
            dokumen_pendukung=e.dokumen_pendukung,
        )
        for e in entries
    ]

@app.post("/api/generate-word")
async def generate_word_document(
    data: Optional[GenerateLogbookData] = None,
    current_user: Optional[User] = Depends(get_current_user_optional),
    db: Session = Depends(get_db),
):
    try:
        # Create a new Document
        doc = Document()
        
        # Set default font to Times New Roman, size 12
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)
        
        # Create table with 4 columns
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        # Let Word auto-fit table to contents/page
        table.autofit = True
        
        # Do not force fixed column widths to avoid overflowing the page
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'NO.'
        hdr_cells[1].text = 'HARI/TGL'
        hdr_cells[2].text = 'JAM'
        hdr_cells[3].text = 'KEGIATAN PER HARI'
        
        # Make header bold and centered
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            # Vertical center header cells
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Prepare entries source: if authenticated, use DB entries of current user (sorted)
        # otherwise, fallback to provided data in request body for backward compatibility
        entries_for_doc: List[GenerateLogbookEntry] = []
        if current_user is not None:
            db_entries = (
                db.query(LogbookEntryORM)
                .filter(LogbookEntryORM.user_id == current_user.id)
                .order_by(LogbookEntryORM.tanggal.asc(), LogbookEntryORM.jam_mulai.asc())
                .all()
            )
            for e in db_entries:
                jam_str = f"{e.jam_mulai.strftime('%H:%M')}"
                if e.jam_selesai:
                    jam_str += f" - {e.jam_selesai.strftime('%H:%M')}"
                entries_for_doc.append(
                    GenerateLogbookEntry(
                        id=str(e.id),
                        tanggal=e.tanggal.strftime('%Y-%m-%d'),
                        jam=jam_str,
                        judul_kegiatan=e.judul_kegiatan,
                        rincian_kegiatan=e.rincian_kegiatan,
                        dokumen_pendukung=e.dokumen_pendukung,
                    )
                )
        elif data and data.entries:
            # sort by tanggal then by jam start
            def sort_key(en: GenerateLogbookEntry):
                try:
                    d = datetime.strptime(en.tanggal, '%Y-%m-%d')
                except Exception:
                    d = datetime.max
                try:
                    start = (en.jam or '').split(' - ')[0].strip()
                    t = datetime.strptime(start, '%H:%M').time()
                except Exception:
                    t = time.max
                return (d, t)

            entries_for_doc = sorted(list(data.entries), key=sort_key)
        else:
            entries_for_doc = []

        # Add data rows
        for idx, entry in enumerate(entries_for_doc, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            # Format date to "DD Month YYYY" in Indonesian
            try:
                dt = datetime.strptime(entry.tanggal, '%Y-%m-%d')
                bulan = [
                    'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
                    'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember'
                ]
                formatted_date = f"{dt.day:02d} {bulan[dt.month - 1]} {dt.year}"
            except Exception:
                formatted_date = entry.tanggal
            row_cells[1].text = formatted_date
            row_cells[2].text = entry.jam
            
            # Format kegiatan cell
            kegiatan_cell = row_cells[3]
            kegiatan_paragraph = kegiatan_cell.paragraphs[0]
            
            # Add title
            judul_run = kegiatan_paragraph.add_run("Judul Kegiatan:\n")
            judul_run.bold = True
            judul_run.font.name = 'Times New Roman'
            judul_run.font.size = Pt(10)
            kegiatan_paragraph.add_run(f"• {entry.judul_kegiatan}\n\n")
            
            # Add rincian
            rincian_run = kegiatan_paragraph.add_run("Rincian Kegiatan:\n")
            rincian_run.bold = True
            rincian_run.font.name = 'Times New Roman'
            rincian_run.font.size = Pt(10)
            kegiatan_paragraph.add_run(f"• {entry.rincian_kegiatan}\n\n")
            
            # Add dokumen pendukung
            if entry.dokumen_pendukung:
                dokumen_run = kegiatan_paragraph.add_run("Dokumen Pendukung:\n\n")
                dokumen_run.bold = True
                dokumen_run.font.name = 'Times New Roman'
                dokumen_run.font.size = Pt(10)
                
                try:
                    # Decode base64 image (handle multiple possible data URL formats)
                    base64_str = entry.dokumen_pendukung
                    if 'base64,' in base64_str:
                        base64_str = base64_str.split('base64,', 1)[1]
                    elif ',' in base64_str:
                        # Fallback: take part after first comma
                        base64_str = base64_str.split(',', 1)[1]
                    # Remove whitespace/newlines
                    base64_str = base64_str.strip()
                    image_data = base64.b64decode(base64_str)
                    
                    # Create temporary file
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                        tmp_file.write(image_data)
                        tmp_file_path = tmp_file.name
                    
                    # Add image to document
                    kegiatan_paragraph.add_run().add_picture(tmp_file_path, width=Inches(2.5))
                    
                    # Clean up temporary file
                    os.unlink(tmp_file_path)
                    
                except Exception as e:
                    print(f"Error adding image: {e}")
                    kegiatan_paragraph.add_run("[Gambar tidak dapat ditampilkan]")
            
            # Ensure font for all paragraphs/runs in the row is Times New Roman 10 and vertical center
            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)

            # Center align the number column
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        filename = f"logbook_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        temp_dir = tempfile.gettempdir()
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        
        return FileResponse(
            filepath,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")

@app.get("/api/health")
async def health_check():
    return {"status": "healthy"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)