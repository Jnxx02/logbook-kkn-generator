from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import base64
import tempfile
import os

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class LogbookEntry(BaseModel):
    id: str
    tanggal: str
    jam: str
    judul_kegiatan: str
    rincian_kegiatan: str
    dokumen_pendukung: Optional[str] = None


class LogbookData(BaseModel):
    entries: List[LogbookEntry]


@app.post("/api/generate-word")
async def generate_word_document(data: LogbookData):
    try:
        doc = Document()

        # Default font for body
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(10)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = True

        hdr_cells = table.rows[0].cells
        headers = ['NO.', 'HARI/TGL', 'JAM', 'KEGIATAN PER HARI']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        bulan = [
            'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
            'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember'
        ]

        for idx, entry in enumerate(data.entries, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)

            try:
                dt = datetime.strptime(entry.tanggal, '%Y-%m-%d')
                formatted_date = f"{dt.day:02d} {bulan[dt.month - 1]} {dt.year}"
            except Exception:
                formatted_date = entry.tanggal
            row_cells[1].text = formatted_date
            row_cells[2].text = entry.jam

            kegiatan_cell = row_cells[3]
            kegiatan_paragraph = kegiatan_cell.paragraphs[0]
            judul_run = kegiatan_paragraph.add_run("Judul Kegiatan:\n")
            judul_run.bold = True
            judul_run.font.name = 'Times New Roman'
            judul_run.font.size = Pt(10)
            kegiatan_paragraph.add_run(f"• {entry.judul_kegiatan}\n\n")

            rincian_run = kegiatan_paragraph.add_run("Rincian Kegiatan:\n")
            rincian_run.bold = True
            rincian_run.font.name = 'Times New Roman'
            rincian_run.font.size = Pt(10)
            kegiatan_paragraph.add_run(f"• {entry.rincian_kegiatan}\n\n")

            if entry.dokumen_pendukung:
                dokumen_run = kegiatan_paragraph.add_run("Dokumen Pendukung:\n\n")
                dokumen_run.bold = True
                dokumen_run.font.name = 'Times New Roman'
                dokumen_run.font.size = Pt(10)
                try:
                    base64_str = entry.dokumen_pendukung
                    if 'base64,' in base64_str:
                        base64_str = base64_str.split('base64,', 1)[1]
                    elif ',' in base64_str:
                        base64_str = base64_str.split(',', 1)[1]
                    base64_str = base64_str.strip()
                    image_data = base64.b64decode(base64_str)
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                        tmp_file.write(image_data)
                        tmp_file_path = tmp_file.name
                    kegiatan_paragraph.add_run().add_picture(tmp_file_path, width=Inches(2.5))
                    os.unlink(tmp_file_path)
                except Exception as e:
                    print(f"Error adding image: {e}")
                    kegiatan_paragraph.add_run("[Gambar tidak dapat ditampilkan]")

            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)

            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        filename = f"logbook_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        temp_dir = tempfile.gettempdir()
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)

        return FileResponse(
            filepath,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")


@app.get("/api/health")
async def health_check():
    return {"status": "healthy"}


