from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional
import json
import gzip
from datetime import datetime, time
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import base64
import tempfile
import os

app = FastAPI()


# Schemas
class GenerateEntry(BaseModel):
    id: str
    tanggal: str
    jam: str
    judul_kegiatan: str
    rincian_kegiatan: str
    dokumen_pendukung: Optional[str] = None


class GenerateBody(BaseModel):
    entries: List[GenerateEntry]


def _is_base64_data_url(value: Optional[str]) -> bool:
    if not value:
        return False
    # data:image/png;base64,....
    return value.startswith("data:") and "base64," in value

@app.get("/")
@app.get("/api/generate-word")
async def generate_word_document_root():
    return {"status": "generate-word"}


@app.post("/api/generate-word")
async def generate_word_document(request: Request):
    try:
        doc = Document()

        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(10)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = True

        hdr_cells = table.rows[0].cells
        headers = ['NO.', 'HARI/TGL', 'JAM', 'KEGIATAN PER HARI']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        bulan = [
            'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
            'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember'
        ]

        entries_for_doc: List[GenerateEntry] = []
        # Accept raw (optionally gzipped) JSON
        parsed_entries: List[GenerateEntry] = []
        try:
            raw_bytes = await request.body()
            if raw_bytes:
                encoding = (request.headers.get('content-encoding') or '').lower()
                if 'gzip' in encoding:
                    try:
                        raw_bytes = gzip.decompress(raw_bytes)
                    except Exception:
                        raise HTTPException(status_code=400, detail="Failed to decompress gzip body")
                try:
                    body_json = json.loads(raw_bytes.decode('utf-8'))
                except Exception:
                    raise HTTPException(status_code=400, detail="Invalid JSON body")
                entries_list = body_json.get('entries') if isinstance(body_json, dict) else None
                if isinstance(entries_list, list):
                    for item in entries_list:
                        if isinstance(item, dict):
                            parsed_entries.append(GenerateEntry(
                                id=str(item.get('id', '')),
                                tanggal=str(item.get('tanggal', '')),
                                jam=str(item.get('jam', '')),
                                judul_kegiatan=str(item.get('judul_kegiatan', '')),
                                rincian_kegiatan=str(item.get('rincian_kegiatan', '')),
                                dokumen_pendukung=item.get('dokumen_pendukung'),
                            ))
        except HTTPException:
            raise
        except Exception:
            parsed_entries = []

        def sort_key(en: GenerateEntry):
            try:
                d = datetime.strptime(en.tanggal, '%Y-%m-%d')
            except Exception:
                d = datetime.max
            try:
                start = (en.jam or '').split(' - ')[0].strip()
                t = datetime.strptime(start, '%H:%M').time()
            except Exception:
                t = time.max
            return (d, t)
        entries_for_doc = sorted(parsed_entries, key=sort_key)

        for idx, entry in enumerate(entries_for_doc, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)

            try:
                dt = datetime.strptime(entry.tanggal, '%Y-%m-%d')
                formatted_date = f"{dt.day:02d} {bulan[dt.month - 1]} {dt.year}"
            except Exception:
                formatted_date = entry.tanggal
            row_cells[1].text = formatted_date
            row_cells[2].text = entry.jam

            kegiatan_cell = row_cells[3]
            kegiatan_paragraph = kegiatan_cell.paragraphs[0]
            judul_run = kegiatan_paragraph.add_run("Judul Kegiatan:\n")
            judul_run.bold = True
            judul_run.font.name = 'Times New Roman'
            judul_run.font.size = Pt(10)
            kegiatan_paragraph.add_run(f"• {entry.judul_kegiatan}\n\n")

            rincian_run = kegiatan_paragraph.add_run("Rincian Kegiatan:\n")
            rincian_run.bold = True
            rincian_run.font.name = 'Times New Roman'
            rincian_run.font.size = Pt(10)
            kegiatan_paragraph.add_run(f"• {entry.rincian_kegiatan}\n\n")

            if entry.dokumen_pendukung:
                dokumen_run = kegiatan_paragraph.add_run("Dokumen Pendukung:\n\n")
                dokumen_run.bold = True
                dokumen_run.font.name = 'Times New Roman'
                dokumen_run.font.size = Pt(10)
                try:
                    image_data = None
                    if _is_base64_data_url(entry.dokumen_pendukung):
                        raw = entry.dokumen_pendukung.split('base64,', 1)[1] if 'base64,' in entry.dokumen_pendukung else entry.dokumen_pendukung
                        raw = raw.split(',', 1)[1] if ',' in raw else raw
                        raw = raw.strip()
                        image_data = base64.b64decode(raw)
                    else:
                        image_data = download_image_bytes(entry.dokumen_pendukung)
                    if not image_data:
                        raise Exception('No image bytes')
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                        tmp_file.write(image_data)
                        tmp_file_path = tmp_file.name
                    kegiatan_paragraph.add_run().add_picture(tmp_file_path, width=Inches(2.5))
                    os.unlink(tmp_file_path)
                except Exception as e:
                    print(f"Error adding image: {e}")
                    kegiatan_paragraph.add_run("[Gambar tidak dapat ditampilkan]")

            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)

            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        filename = f"logbook_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        temp_dir = tempfile.gettempdir()
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)

        return FileResponse(
            filepath,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")


